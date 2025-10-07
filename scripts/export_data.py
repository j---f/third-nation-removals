import json
import csv
import pandas as pd
from datetime import datetime
import os

def load_removals_data():
    """Load removals data from JSON file"""
    try:
        with open('data/removals.json', 'r') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def export_to_json(data, filename=None):
    """Export data to JSON format"""
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    
    os.makedirs('exports', exist_ok=True)
    with open(filename, 'w') as f:
        json.dump(data, f, indent=2)
    
    return filename

def export_to_csv(data, filename=None):
    """Export data to CSV format"""
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    
    os.makedirs('exports', exist_ok=True)
    
    # Flatten the data for CSV export
    flattened_data = []
    for entry in data:
        flat_entry = {
            'destination_country': entry.get('destination_country', ''),
            'date': entry.get('date', ''),
            'date_range_end': entry.get('date_range_end', ''),
            'number_removed': entry.get('number_removed', ''),
            'origin_nationalities': ', '.join(entry.get('origin_nationalities', [])),
            'source_urls': ', '.join(entry.get('source_urls', [])),
            'notes': entry.get('notes', '')
        }
        flattened_data.append(flat_entry)
    
    with open(filename, 'w', newline='') as f:
        if flattened_data:
            writer = csv.DictWriter(f, fieldnames=flattened_data[0].keys())
            writer.writeheader()
            writer.writerows(flattened_data)
    
    return filename

def export_to_md(data, filename=None):
    """Export data to Markdown format"""
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    
    os.makedirs('exports', exist_ok=True)
    
    with open(filename, 'w') as f:
        f.write("# Third-Nation Removals Data\n\n")
        f.write(f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n")
        
        # Summary statistics
        total_removals = len(data)
        total_people = sum(entry.get('number_removed', 0) for entry in data if entry.get('number_removed'))
        countries = list(set(entry.get('destination_country', '') for entry in data))
        
        f.write("## Summary\n\n")
        f.write(f"- Total removal events: {total_removals}\n")
        f.write(f"- Total people removed: {total_people}\n")
        f.write(f"- Destination countries: {len(countries)}\n\n")
        
        # Country breakdown
        f.write("## By Destination Country\n\n")
        country_data = {}
        for entry in data:
            country = entry.get('destination_country', '')
            if country not in country_data:
                country_data[country] = {
                    'events': 0,
                    'people': 0
                }
            country_data[country]['events'] += 1
            if entry.get('number_removed'):
                country_data[country]['people'] += entry['number_removed']
        
        for country, stats in sorted(country_data.items()):
            f.write(f"### {country}\n")
            f.write(f"- Events: {stats['events']}\n")
            f.write(f"- People removed: {stats['people']}\n\n")
        
        # Detailed data
        f.write("## Detailed Data\n\n")
        f.write("| Destination Country | Date | Date Range End | Number Removed | Origin Nationalities | Notes |\n")
        f.write("|---|---|---|---|---|---|\n")
        
        for entry in data:
            destination = entry.get('destination_country', '')
            date = entry.get('date', '')
            date_end = entry.get('date_range_end', '')
            number = entry.get('number_removed', '')
            nationalities = ', '.join(entry.get('origin_nationalities', []))
            notes = entry.get('notes', '').replace('\n', ' ')[:100] + ('...' if len(entry.get('notes', '')) > 100 else '')
            
            f.write(f"| {destination} | {date} | {date_end} | {number} | {nationalities} | {notes} |\n")
    
    return filename

def export_to_txt(data, filename=None):
    """Export data to plain text format"""
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    
    os.makedirs('exports', exist_ok=True)
    
    with open(filename, 'w') as f:
        f.write("THIRD-NATION REMOVALS DATA\n")
        f.write("=" * 50 + "\n")
        f.write(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        # Summary statistics
        total_removals = len(data)
        total_people = sum(entry.get('number_removed', 0) for entry in data if entry.get('number_removed'))
        countries = list(set(entry.get('destination_country', '') for entry in data))
        
        f.write("SUMMARY\n")
        f.write("-" * 20 + "\n")
        f.write(f"Total removal events: {total_removals}\n")
        f.write(f"Total people removed: {total_people}\n")
        f.write(f"Destination countries: {len(countries)}\n\n")
        
        # Country breakdown
        f.write("BY DESTINATION COUNTRY\n")
        f.write("-" * 30 + "\n")
        country_data = {}
        for entry in data:
            country = entry.get('destination_country', '')
            if country not in country_data:
                country_data[country] = {
                    'events': 0,
                    'people': 0
                }
            country_data[country]['events'] += 1
            if entry.get('number_removed'):
                country_data[country]['people'] += entry['number_removed']
        
        for country, stats in sorted(country_data.items()):
            f.write(f"{country}\n")
            f.write(f"  Events: {stats['events']}\n")
            f.write(f"  People removed: {stats['people']}\n\n")
        
        # Detailed data
        f.write("DETAILED DATA\n")
        f.write("-" * 20 + "\n")
        
        for i, entry in enumerate(data, 1):
            f.write(f"Entry {i}\n")
            f.write(f"  Destination Country: {entry.get('destination_country', '')}\n")
            f.write(f"  Date: {entry.get('date', '')}\n")
            f.write(f"  Date Range End: {entry.get('date_range_end', '')}\n")
            f.write(f"  Number Removed: {entry.get('number_removed', '')}\n")
            f.write(f"  Origin Nationalities: {', '.join(entry.get('origin_nationalities', []))}\n")
            f.write(f"  Source URLs: {', '.join(entry.get('source_urls', []))}\n")
            f.write(f"  Notes: {entry.get('notes', '')}\n\n")
    
    return filename

def export_to_pdf(data, filename=None):
    """Export data to PDF format"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        from reportlab.lib.units import inch
    except ImportError:
        print("Error: reportlab is required for PDF export. Install with: pip install reportlab")
        return None
    
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    os.makedirs('exports', exist_ok=True)
    
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph("Third-Nation Removals Data", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Generation date
    gen_date = Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
    story.append(gen_date)
    story.append(Spacer(1, 12))
    
    # Summary statistics
    total_removals = len(data)
    total_people = sum(entry.get('number_removed', 0) for entry in data if entry.get('number_removed'))
    countries = list(set(entry.get('destination_country', '') for entry in data))
    
    summary_title = Paragraph("Summary", styles['Heading2'])
    story.append(summary_title)
    story.append(Spacer(1, 6))
    
    summary_data = [
        ['Metric', 'Value'],
        ['Total removal events', str(total_removals)],
        ['Total people removed', str(total_people)],
        ['Destination countries', str(len(countries))]
    ]
    
    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(summary_table)
    story.append(Spacer(1, 12))
    
    # Country breakdown
    country_title = Paragraph("By Destination Country", styles['Heading2'])
    story.append(country_title)
    story.append(Spacer(1, 6))
    
    country_data = {}
    for entry in data:
        country = entry.get('destination_country', '')
        if country not in country_data:
            country_data[country] = {
                'events': 0,
                'people': 0
            }
        country_data[country]['events'] += 1
        if entry.get('number_removed'):
            country_data[country]['people'] += entry['number_removed']
    
    country_table_data = [['Country', 'Events', 'People Removed']]
    for country, stats in sorted(country_data.items()):
        country_table_data.append([country, str(stats['events']), str(stats['people'])])
    
    country_table = Table(country_table_data)
    country_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(country_table)
    story.append(Spacer(1, 12))
    
    # Detailed data (first 20 entries to keep PDF manageable)
    detailed_title = Paragraph("Detailed Data (First 20 Entries)", styles['Heading2'])
    story.append(detailed_title)
    story.append(Spacer(1, 6))
    
    detailed_table_data = [
        ['Country', 'Date', 'Number', 'Nationalities', 'Notes (Truncated)']
    ]
    
    for entry in data[:20]:
        destination = entry.get('destination_country', '')
        date = entry.get('date', '')
        number = str(entry.get('number_removed', ''))
        nationalities = ', '.join(entry.get('origin_nationalities', []))
        notes = entry.get('notes', '').replace('\n', ' ')[:50] + ('...' if len(entry.get('notes', '')) > 50 else '')
        
        detailed_table_data.append([destination, date, number, nationalities, notes])
    
    detailed_table = Table(detailed_table_data, colWidths=[1.2*inch, 0.8*inch, 0.6*inch, 1.2*inch, 2*inch])
    detailed_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(detailed_table)
    
    # Build PDF
    doc.build(story)
    
    return filename

def export_to_doc(data, filename=None):
    """Export data to Word document format"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("Error: python-docx is required for DOC export. Install with: pip install python-docx")
        return None
    
    if filename is None:
        filename = f"exports/third_nation_removals_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    os.makedirs('exports', exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Add title
    doc.add_heading('Third-Nation Removals Data', 0)
    doc.add_paragraph(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Summary statistics
    total_removals = len(data)
    total_people = sum(entry.get('number_removed', 0) for entry in data if entry.get('number_removed'))
    countries = list(set(entry.get('destination_country', '') for entry in data))
    
    doc.add_heading('Summary', level=1)
    p = doc.add_paragraph()
    p.add_run('Total removal events: ').bold = True
    p.add_run(f'{total_removals}\n')
    p.add_run('Total people removed: ').bold = True
    p.add_run(f'{total_people}\n')
    p.add_run('Destination countries: ').bold = True
    p.add_run(f'{len(countries)}')
    
    # Country breakdown
    doc.add_heading('By Destination Country', level=1)
    
    country_data = {}
    for entry in data:
        country = entry.get('destination_country', '')
        if country not in country_data:
            country_data[country] = {
                'events': 0,
                'people': 0
            }
        country_data[country]['events'] += 1
        if entry.get('number_removed'):
            country_data[country]['people'] += entry['number_removed']
    
    # Add table for country data
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Country'
    hdr_cells[1].text = 'Events'
    hdr_cells[2].text = 'People Removed'
    
    for country, stats in sorted(country_data.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = country
        row_cells[1].text = str(stats['events'])
        row_cells[2].text = str(stats['people'])
    
    # Detailed data
    doc.add_heading('Detailed Data', level=1)
    
    # Add table for detailed data
    detailed_table = doc.add_table(rows=1, cols=6)
    detailed_table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = detailed_table.rows[0].cells
    hdr_cells[0].text = 'Destination Country'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Date Range End'
    hdr_cells[3].text = 'Number Removed'
    hdr_cells[4].text = 'Origin Nationalities'
    hdr_cells[5].text = 'Notes'
    
    for entry in data:
        row_cells = detailed_table.add_row().cells
        row_cells[0].text = entry.get('destination_country', '')
        row_cells[1].text = entry.get('date', '')
        row_cells[2].text = entry.get('date_range_end', '')
        row_cells[3].text = str(entry.get('number_removed', ''))
        row_cells[4].text = ', '.join(entry.get('origin_nationalities', []))
        row_cells[5].text = entry.get('notes', '')
    
    # Save document
    doc.save(filename)
    
    return filename

def main():
    """Main function to export data in various formats"""
    data = load_removals_data()
    
    if not data:
        print("No data found to export.")
        return
    
    print(f"Loaded {len(data)} removal records.")
    
    # Export to all formats
    formats = [
        ('json', export_to_json),
        ('csv', export_to_csv),
        ('md', export_to_md),
        ('txt', export_to_txt),
        ('pdf', export_to_pdf),
        ('doc', export_to_doc)
    ]
    
    exported_files = []
    
    for format_name, export_func in formats:
        try:
            filename = export_func(data)
            if filename:
                exported_files.append(filename)
                print(f"Exported to {format_name.upper()}: {filename}")
        except Exception as e:
            print(f"Error exporting to {format_name.upper()}: {e}")
    
    print(f"\nExport complete. {len(exported_files)} files created.")
    return exported_files

if __name__ == "__main__":
    main()